#!/usr/bin/env python3
"""
report_builder.py

Assemble a legal review report in Microsoft Word format based on the LLM
clause analysis and deterministic checks. The report contains a cover
page, an executive summary, a clause verdict matrix, and snapshots of
deterministic checks. It also exports a CSV and JSON issue register for
machine consumption.

Usage:
    python report_builder.py \
        --results judged_results.json \
        --deterministic deterministic_results.json \
        --output report.docx \
        --issues-csv issues.csv \
        --issues-json issues.json

The input `judged_results.json` should contain a dictionary mapping
clause IDs to objects with at least the keys: `expected`, `actual`,
`status`, `fix`, `judge_status` (and optionally `severity`). The
`deterministic_results.json` should contain key/value pairs for each
deterministic check with nested fields `status`, etc.
"""
import argparse
import csv
import json
from typing import Dict, Any, List

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    Document = None  # type: ignore
    RGBColor = None  # type: ignore
    Pt = None  # type: ignore

try:
    import diff_match_patch
except ImportError:
    diff_match_patch = None  # type: ignore


def load_json(path: str) -> Any:
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def build_issue_register(results: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Flatten the clause results into a list of issue dicts."""
    issues = []
    for clause_id, data in results.items():
        issues.append({
            'clause': clause_id,
            'expected': data.get('expected', ''),
            'actual': data.get('actual', ''),
            'status': data.get('status', ''),
            'judge_status': data.get('judge_status', ''),
            'fix': data.get('fix', ''),
            'severity': data.get('severity', '')
        })
    return issues


def write_csv(issues: List[Dict[str, Any]], path: str) -> None:
    fieldnames = ['clause', 'status', 'judge_status', 'severity', 'expected', 'actual', 'fix']
    with open(path, 'w', encoding='utf-8', newline='') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for issue in issues:
            writer.writerow(issue)


def write_json(issues: List[Dict[str, Any]], path: str) -> None:
    with open(path, 'w', encoding='utf-8') as f_out:
        json.dump(issues, f_out, ensure_ascii=False, indent=2)


def diff_text(expected: str, actual: str) -> List[tuple]:
    """Compute a list of diffs between expected and actual using diff-match-patch."""
    if diff_match_patch is None:
        return []
    dmp = diff_match_patch.diff_match_patch()
    diffs = dmp.diff_main(expected, actual)
    dmp.diff_cleanupSemantic(diffs)
    return diffs


def add_diff_to_doc(paragraph, diffs) -> None:
    """Add colour-coded diffs to a Word paragraph."""
    if Document is None or RGBColor is None:
        return
    for op, text in diffs:
        run = paragraph.add_run(text)
        if op == diff_match_patch.diff_delete:
            # deletion – red strike
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.strike = True
        elif op == diff_match_patch.diff_insert:
            # insertion – green bold
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            run.font.bold = True
        else:
            # equal – default formatting
            pass


def build_report(results: Dict[str, Any], deterministic: Dict[str, Any], output_path: str, issues: List[Dict[str, Any]]) -> None:
    if Document is None or Pt is None:
        raise RuntimeError("python-docx is not installed")
    doc = Document()
    # Cover page
    doc.add_heading('Dana Energy PO – Legal Review Report', level=1)
    doc.add_paragraph('This report summarises the legal review of the provided Purchase Order package.')

    # Executive summary
    doc.add_heading('Executive Summary', level=2)
    fails = [issue for issue in issues if issue['status'] == 'FAIL' or issue['judge_status'] in ('UNCERTAIN', 'CONFLICT')]
    if fails:
        doc.add_paragraph('The following clauses require attention:')
        for issue in fails:
            doc.add_paragraph(f"• {issue['clause']} – {issue['status']} / {issue['judge_status']}")
    else:
        doc.add_paragraph('All mandatory clauses appear compliant based on the current checks.')

    # Clause verdict matrix
    doc.add_heading('Clause Verdict Matrix', level=2)
    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Clause'
    hdr_cells[1].text = 'Expected'
    hdr_cells[2].text = 'Actual'
    hdr_cells[3].text = 'Status'
    hdr_cells[4].text = 'Judge'
    hdr_cells[5].text = 'Fix'
    for issue in issues:
        row_cells = table.add_row().cells
        row_cells[0].text = issue['clause']
        row_cells[1].text = issue['expected']
        row_cells[2].text = issue['actual']
        row_cells[3].text = issue['status']
        row_cells[4].text = issue['judge_status']
        row_cells[5].text = issue['fix']

    # Deterministic checks snapshot
    doc.add_heading('Deterministic Checks', level=2)
    for key, value in deterministic.items():
        status = value.get('status', '')
        doc.add_paragraph(f"{key}: {status}")

    # Annex A: Diff details per clause
    doc.add_heading('Annex A – Clause Diffs', level=2)
    for clause_id, data in results.items():
        doc.add_heading(clause_id, level=3)
        expected = data.get('expected', '')
        actual = data.get('actual', '')
        diffs = diff_text(expected, actual)
        para = doc.add_paragraph()
        if diffs:
            add_diff_to_doc(para, diffs)
        else:
            para.add_run('No differences detected or diff-match-patch unavailable.')

    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build legal review report")
    parser.add_argument("--results", required=True, help="Path to judged results JSON")
    parser.add_argument("--deterministic", required=True, help="Path to deterministic checks JSON")
    parser.add_argument("--output", required=True, help="Path to output DOCX report")
    parser.add_argument("--issues-csv", required=True, help="Path to write issues CSV")
    parser.add_argument("--issues-json", required=True, help="Path to write issues JSON")
    args = parser.parse_args()

    results = load_json(args.results)
    deterministic = load_json(args.deterministic)
    issues = build_issue_register(results)
    write_csv(issues, args.issues_csv)
    write_json(issues, args.issues_json)
    build_report(results, deterministic, args.output, issues)

if __name__ == '__main__':
    main()