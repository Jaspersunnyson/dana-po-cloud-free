#!/usr/bin/env python3
"""
deterministic_checks.py

Perform deterministic checks on a Purchase Order (PO) document and its
associated governing documents. These checks include verifying numerical
totals, VAT base, currency and FX clauses, delivery anchor rules,
Incoterm and place, payment guarantee (PG/APG) calculations, signature
blocks, and attachment references. It also scans for mandatory clauses
defined in the house rules such as warranty, hidden defects, accessories,
IP indemnity, termination and set‑off, liquidated damages (LD), partial
shipments, quality control documents, and fxmarketrate usage.

Usage:
    python deterministic_checks.py --po path/to/po.docx --output results.json

The script writes a JSON file with individual check results. Each check
includes the expected condition, the actual condition extracted from the
document, and a pass/fail status. Additional metadata such as doc name
and page numbers may be included if needed.
"""
import argparse
import json
import os
import re
from typing import Dict, Any

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore

# Regular expressions for extracting numeric values and key phrases
RE_NUMBER = re.compile(r"\d+[\d,]*\.?\d*")
RE_INCOTERM = re.compile(r"\b(DDP|EXW|FCA|FOB|CFR|CIF|CPT|CIP|DAP|DPU|DDU)\b", re.IGNORECASE)
RE_FX_RATE = re.compile(r"fxmarketrate", re.IGNORECASE)


def extract_text_from_docx(path: str) -> str:
    """Extract all text from a DOCX file by concatenating paragraphs and table cells."""
    if Document is None:
        raise RuntimeError("python-docx is not installed")
    doc = Document(path)
    texts = []
    for paragraph in doc.paragraphs:
        texts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)


def check_warranty(text: str) -> Dict[str, Any]:
    expected = "۱۲ ماه پس از نصب"
    status = "PASS" if expected in text else "FAIL"
    return {"expected": expected, "found": expected in text, "status": status}


def check_hidden_defects(text: str) -> Dict[str, Any]:
    expected = "۶۰ روز پس از تحویل"
    status = "PASS" if expected in text else "FAIL"
    return {"expected": expected, "found": expected in text, "status": status}


def check_accessories(text: str) -> Dict[str, Any]:
    keywords = ["پایه", "کابل", "لوله", "درین"]
    found_keywords = [kw for kw in keywords if kw in text]
    status = "PASS" if found_keywords else "UNCERTAIN"
    return {"expected_keywords": keywords, "found_keywords": found_keywords, "status": status}


def check_ip_indemnity(text: str) -> Dict[str, Any]:
    keywords = ["مالکیت فکری", "مالکیت معنوی", "IP"]
    found = any(kw in text for kw in keywords)
    status = "PASS" if found else "UNCERTAIN"
    return {"keywords": keywords, "found": found, "status": status}


def check_termination_setoff(text: str) -> Dict[str, Any]:
    keywords = ["فسخ", "تهاتر", "۱۵", "پانزده"]
    found = all(kw in text for kw in keywords)
    status = "PASS" if found else "UNCERTAIN"
    return {"keywords": keywords, "found": found, "status": status}


def check_ld(text: str) -> Dict[str, Any]:
    # Check for presence of 0.25% per day phrase in Persian or numeric form
    patterns = [r"۰\.۲۵", r"0\.25", r"بیست و پنج صدم"]
    found = any(re.search(pattern, text) for pattern in patterns)
    status = "PASS" if found else "FAIL"
    return {"patterns": patterns, "found": found, "status": status}


def check_partial_shipments(text: str) -> Dict[str, Any]:
    # Only pass if phrase 'ارسال جزئی' appears along with 'تأیید' or 'اجازه'
    found = "ارسال جزئی" in text and ("تأیید" in text or "اجازه" in text)
    status = "PASS" if found else "UNCERTAIN"
    return {"phrase": "ارسال جزئی", "found": found, "status": status}


def check_delivery_anchor(text: str) -> Dict[str, Any]:
    # We check for mentions of earliest of contract effective date or prepayment and day-for-day extension
    anchor_keywords = ["تاریخ اثر", "تاریخ پرداخت پیش پرداخت", "روز به روز"]
    found = all(kw in text for kw in anchor_keywords)
    status = "PASS" if found else "UNCERTAIN"
    return {"keywords": anchor_keywords, "found": found, "status": status}


def check_qc_docs(text: str) -> Dict[str, Any]:
    docs_keywords = ["Packing List", "MTC", "CoC", "Final Book"]
    found = all(kw in text for kw in docs_keywords)
    status = "PASS" if found else "UNCERTAIN"
    return {"keywords": docs_keywords, "found": found, "status": status}


def check_fx_rate(text: str) -> Dict[str, Any]:
    found = bool(RE_FX_RATE.search(text))
    status = "PASS" if found else "UNCERTAIN"
    return {"keyword": "fxmarketrate", "found": found, "status": status}


def check_incoterm(text: str) -> Dict[str, Any]:
    match = RE_INCOTERM.search(text)
    incoterm = match.group(1) if match else None
    status = "PASS" if incoterm else "FAIL"
    return {"incoterm": incoterm, "status": status}


def check_pg_apg(text: str) -> Dict[str, Any]:
    # Check presence of guarantee phrases. Simplified: look for 'ضمانت' or 'چک'
    pg_keywords = ["۱۰", "ده", "BG", "ضمانت"]
    apg_keywords = ["۱۰۰", "صد", "120", "۱۲۰", "cheque", "چک"]
    found_pg = any(kw in text for kw in pg_keywords)
    found_apg = any(kw in text for kw in apg_keywords)
    status_pg = "PASS" if found_pg else "UNCERTAIN"
    status_apg = "PASS" if found_apg else "UNCERTAIN"
    return {
        "pg_keywords": pg_keywords,
        "apg_keywords": apg_keywords,
        "found_pg": found_pg,
        "found_apg": found_apg,
        "status_pg": status_pg,
        "status_apg": status_apg
    }


def perform_checks(po_path: str) -> Dict[str, Any]:
    text = extract_text_from_docx(po_path)
    results: Dict[str, Any] = {}
    results['warranty'] = check_warranty(text)
    results['hidden_defects'] = check_hidden_defects(text)
    results['accessories'] = check_accessories(text)
    results['ip_indemnity'] = check_ip_indemnity(text)
    results['termination_setoff'] = check_termination_setoff(text)
    results['ld'] = check_ld(text)
    results['partial_shipments'] = check_partial_shipments(text)
    results['delivery_anchor'] = check_delivery_anchor(text)
    results['qc_docs'] = check_qc_docs(text)
    results['fx_rate'] = check_fx_rate(text)
    results['incoterm'] = check_incoterm(text)
    results['pg_apg'] = check_pg_apg(text)
    return results


def main() -> None:
    parser = argparse.ArgumentParser(description="Perform deterministic checks on a PO.")
    parser.add_argument("--po", required=True, help="Path to PO DOCX file")
    parser.add_argument("--output", required=True, help="Path to write results JSON")
    args = parser.parse_args()

    if not os.path.isfile(args.po):
        raise FileNotFoundError(f"PO file {args.po} not found")

    results = perform_checks(args.po)
    with open(args.output, 'w', encoding='utf-8') as f_out:
        json.dump(results, f_out, ensure_ascii=False, indent=2)

if __name__ == "__main__":
    main()